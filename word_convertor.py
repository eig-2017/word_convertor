# -*- coding: utf-8 -*-
import re
import pandas as pd
from docx import Document

document = Document('examples/RPT-ETAT-ACTIONNAIRE.docx')

# Get the structure of the docx file text
#text_df = pd.DataFrame(data = [(para.text, para.style.name) for para in document.paragraphs], columns = ("Text", "Style"))    
#print("They are %s different text styles in the document." %text_df["Style"].nunique())

def table_to_html(table):
    """
    Get a docx table, generate a html table
    """
    html_table = "<table class = \"table\">"
    for row in table.rows:
        html_table += "<tr>"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                html_table += "<td>" + paragraph.text + "</td>"
        html_table += "</tr>"
    html_table += "</table>"
    return(html_table)
                
# Get the structure of the docx file table
#table_df = pd.DataFrame(data = [(table_to_html(table), table.style.name) for table in document.tables], columns = ("Table","Style"))
#print("They are %s different text styles in the document." %table_df["Style"].nunique())

from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def text_or_table(block):
    """
    If the block is a paragraph, return text.
    If the block is a table, return an html table
    """
    try:
        return(block.text)
    except AttributeError:
        return(table_to_html(block))

# Build a dataframe containing either the text (for paragraph) or content (for tables), and the style
structure = pd.DataFrame(data = [(text_or_table(block),block.style.name) for block in iter_block_items(document)], columns = ["content", "style"])

# A dictionnary between Word styles and HTML tags
style_correspondance = {
"Titre général" : "h1",
"Heading 1" : "h1",
"Heading 2" : "h2",
"Heading 3" : "h3",
"Heading 4" : "h4",
"Heading 5" : "h5",
"Heading 6" : "h6",
"Titre de partie" : "h1",
"Titre Annexe" : "h1",
"Titre (Sommaire des réponses)" : "h1",
"Réponse" : "i",
"Puce Reco + Italique" : "i",
"ENCADRÉ - Titre" : "ENCADRÉ - Texte",
}
default_tag = "p"

# Create a new column with HTML tags
structure["HTML"] = structure["style"].map(style_correspondance).fillna(default_tag)

# Create a new column with HTML syntax, tags and content
structure["web content"] = "<" + structure["HTML"] + ">" + structure["content"] + "</" + structure["HTML"] + ">"
          
# Rename styles (used by Word), to classes (used by CSS)
style = pd.DataFrame(data = structure["style"].unique(), columns = ["style"])
style["class"] = [re.sub('[-+()/°]', ' ', text) for text in style["style"]] #Remove special characters
style["class"] = [re.sub(r"\s+", '-', text) for text in style["class"]] # Replace spaces with dash
style["class"] = [text.lower() for text in style["class"]] # Lower
style = dict(zip(style["style"],style["class"]))
structure["class"] = structure["style"].map(style)     

# Add class tag
structure["web content"] = "<div class =\"" + structure["class"] + "\">" + structure["web content"] + "</div>"

# HTML header
header = """
<!DOCTYPE html PUBLIC "-//W3C//DTD HTML 4.01//EN"
    "http://www.w3.org/TR/html4/strict.dtd">
<html lang="en">
  <head>
    <meta http-equiv="content-type" content="text/html; charset=utf-8">
    <title>C2C</title>
    <link rel="stylesheet" href="style.css">
    <!-- Bootstrap– latest compiled and minified CSS -->
    <link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/3.3.7/css/bootstrap.min.css" integrity="sha384-BVYiiSIFeK1dGmJRAkycuHAHRg32OmUcww7on3RYdg4Va+PmSTsz/K68vbdEjh4u" crossorigin="anonymous">
  </head>
  <body>
  <div class = "container">
"""

# HTML footer
footer = """
    </div>
</body>
</html>
"""

# Heater + content + footer
webpage = pd.Series(header).append(structure["web content"], ignore_index = True)
webpage = webpage.append(pd.Series(footer), ignore_index = True)

# Write the webpage
with open('index.html', 'w') as myFile:
    for line in webpage:
        myFile.write(line + "\n")